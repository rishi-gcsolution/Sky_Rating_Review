# -*- coding: utf-8 -*-
"""
Created on Fri Nov 10 17:28:44 2023

@author: asus
"""

# -*- coding: utf-8 -*-
"""
Created on Fri Nov 10 11:35:03 2023

@author: asus
"""
import os
import pandas as pd
import requests
from datetime import datetime
from google_play_scraper import app, Sort, reviews
from datetime import datetime,timedelta
import matplotlib.pyplot as plt
from docx import Document

# Replace 'com.cloudtradetech.sky' with the package name of the app you want to scrape.
app_id = 'com.cloudtradetech.sky'

# Specify language and country
language = 'en'
country = 'in'

 
# Specify start and end dates for filtering format date= yyyy-mm-dd
start_date_str = "2023-11-01"  # Replace with your desired start date
end_date_str = "2023-11-03"  # Replace with your desired end date



def get_current_datetime():
    try:
        response = requests.get("http://worldtimeapi.org/api/ip")
        response.raise_for_status()  # Raise an exception for bad responses

        data = response.json()
        current_datetime_str = data['datetime']
        # Convert the datetime string to a datetime object
        current_datetime = datetime.strptime(current_datetime_str, "%Y-%m-%dT%H:%M:%S.%f%z")

        return current_datetime

    except requests.exceptions.RequestException as e:
        print(f"Error fetching current datetime: {e}")
        return None
   

def full_filtered_review(start_date_str,end_date_str):
    
    #Convert start and end dates to datetime objects
    start_date = datetime.strptime(start_date_str, "%Y-%m-%d")
    end_date = datetime.strptime(end_date_str, "%Y-%m-%d")
    
    # Fetch the reviews sorted by date in descending order
    reviews_result, _ = reviews(
        app_id,
        lang=language,
        country=country,
        sort=Sort.NEWEST,
        count=10000,  # Number of reviews to retrieve (adjust as needed)
    )
    
    # Filter reviews based on date range given as start date and end date
    filtered_reviews = [
        review
        for review in reviews_result
        if start_date <= review['at'] <= end_date
    ]
    return (reviews_result,filtered_reviews)

##################################
# Current datetime module
##################################

 
    
    
  
#fetching the current datetime
current_datetime = get_current_datetime()

if current_datetime:
    print(f"Current Datetime: {current_datetime}")
else:
    print("Failed to fetch current datetime.")

############################

################################
#calculation of overall rating and separate rating based on  datetime interval

full_filtered=full_filtered_review(start_date_str,end_date_str)

overall_rating=0
interval_rating=0
for review in full_filtered[0]:
    overall_rating+=review['score']
overall_rating=overall_rating/len(full_filtered[0])

interval_rating = sum([
    review['score']
    for review in full_filtered[1]])/len(full_filtered[1])



###################################
#Enter the no of T minus days
T_minus_days=10

# Subtract days from the current datetime
#result_datetime = current_datetime - timedelta(days=T_minus_days)
#######################################################################################################

#######################################################################################################


def Trend_of_Tminus_days(T_minus_days):
    
    
    today_date=get_current_datetime().strftime("%Y-%m-%d")
    end_date=(current_datetime - timedelta(days=T_minus_days)).strftime("%Y-%m-%d")
    #full_filtered_review() takes string type date 
    trend_line_data=full_filtered_review(end_date,today_date)
    lis=[]
    required_date=(current_datetime - timedelta(days=T_minus_days)).strftime("%Y-%m-%d")
    def scoreSum_and_date(required_date):
        
        day_sum=[score['score'] for score in trend_line_data[0]
                 if score['at'].strftime("%Y-%m-%d")==required_date
                 ]
        return [sum(day_sum),required_date,len(day_sum)]
        
    while T_minus_days>0:
        required_date=(current_datetime - timedelta(days=T_minus_days)).strftime("%Y-%m-%d")
        lis.append(scoreSum_and_date(required_date))
        T_minus_days-=1
    lis = [
    [sublist[0] / sublist[2],sublist[1]] if sublist[2] != 0 else [0,sublist[1]]
    for sublist in lis
    ]
    return lis       
    

xx=Trend_of_Tminus_days(T_minus_days)
# Convert list of lists to DataFrame
df_xx = pd.DataFrame(xx, columns=['avg daywise rating', 'Date'])
df_xx = df_xx[['Date', 'avg daywise rating']]


# Extract x and y values from the list
x_values = [pair[1] for pair in xx]
y_values = [pair[0] for pair in xx]

# Plot the line plot
plt.figure(figsize=(10, 5))
plt.subplot(1, 2, 1)  # 1 row, 2 columns, subplot 1
plt.plot(x_values, y_values, marker='o', linestyle='-', color='b')
plt.xlabel('X-axis')
plt.ylabel('Y-axis')
plt.title('Line Plot: xx[0] on Y-axis and xx[1] on X-axis')
plt.grid(True)

# Create a histogram
plt.subplot(1, 2, 2)  # 1 row, 2 columns, subplot 2
plt.hist2d(x_values, y_values, bins=(10, 10), cmap=plt.cm.Blues)
plt.xlabel('X Axis')
plt.ylabel('Y Axis')
plt.title('Histogram of X and Y values')
plt.colorbar()

# Show the plot
plt.tight_layout()  # Adjust layout to prevent overlapping
plt.show()



#######################################################
# Split based on total reviews given
time_duration=11
#T_minus days is time duration
def net_split_of_all_reviews(time_duration):
    net_split={'5':0,'4':0,'3':0,'2':0,'1':0}
    if time_duration=="all":
        for score in full_filtered[0]:
            net_split[str(score['score'])]+=1
            
    if isinstance(time_duration, int):    
        for score in full_filtered[1]:
            net_split[str(score['score'])]+=1
    
    return net_split
            

net_split_data = net_split_of_all_reviews(time_duration)

# Plotting
plt.bar(net_split_data.keys(), net_split_data.values())
plt.xlabel('Score')
plt.ylabel('Count')
plt.title('Net Split of All Reviews')
plt.show()
           
#########################################################################
#split trend wise with respect to previous days this function is giving no of 5star 4star 3star 2star 1star given perticular days
########################################################################

def split_trendwise_wrt_previous_days(T_minus_days):
    def rating_count_fun(data_input):
        #data_input=full_filtered[0]
        # Create a DataFrame
        df = pd.DataFrame(data_input)
        #converting datetime to string of format yyyy-mm-dd
        df['at'] = df['at'].dt.strftime("%Y-%m-%d")
       
            
        # Count the number of ratings for each score on each day
        rating_counts = df.groupby(['at', 'score']).size().unstack(fill_value=0)
            
        # Rename columns for clarity
        rating_counts.columns = [f'{col} Star Ratings' for col in rating_counts.columns]
            
        # Add a column for Total Ratings
        rating_counts['Total Ratings'] = rating_counts.sum(axis=1)
        
        # Reset the index for a clean DataFrame
        rating_counts.reset_index(inplace=True)
        
        return rating_counts
    
    if T_minus_days=="all":
        rating_count_fun(full_filtered[0])
    if isinstance(T_minus_days, int):
        rating_count_fun(full_filtered[1])
    return 0
    
tttttt=split_trendwise_wrt_previous_days(T_minus_days)
    
def thumbs_count():
    pass

def get_no_Of_ratingsYesterday_avgRating():
    yesterday=(get_current_datetime() - timedelta(days=2)).strftime("%Y-%m-%d")
    yesterdays_review=full_filtered_review(yesterday,get_current_datetime().strftime("%Y-%m-%d"))
    net_split_of_all_reviews()

###########################################################################
# saving to ms Word

# Specify the file path for Word document
docx_file_path = r"C:\Users\asus\Desktop\HDFC Work Related\document.docx"


# Create a Word document
doc = Document()

# Write the date to the document

data_to_write = f"                  REVIEW AND RATING ANALYSIS OF SKY\n"
doc.add_paragraph(data_to_write)

data_to_write = f" Date of Analysis = {current_datetime.strftime('%Y-%m-%d')}\n also LAST {T_minus_days} analysis \n"
doc.add_paragraph(data_to_write)

# Write two lines of empty space
doc.add_paragraph("\n" * 2)

data_to_write = f"Over All Rating of all Reviews given = {overall_rating}\n also Given LAST {T_minus_days} interval Rating = {interval_rating} out {len(full_filtered[1])} given reviews\n"
doc.add_paragraph(data_to_write)

# Write two lines of empty space
doc.add_paragraph("\n" * 2)


data_to_write = f"Rating Avg. of Yesterday = {} also No of Rating give = {} n"
doc.add_paragraph(data_to_write)



# Write the DataFrame to the document in tabular form
doc.add_paragraph("DataFrame:")
table = doc.add_table(df_xx.shape[0]+1, df_xx.shape[1])
for col_num, col_name in enumerate(df_xx.columns):
    table.cell(0, col_num).text = col_name
    for row_num in range(df_xx.shape[0]):
        table.cell(row_num+1, col_num).text = str(df_xx.iloc[row_num, col_num])



# Write four lines of empty space
doc.add_paragraph("\n" * 4)

# Write net split data to the document
doc.add_paragraph("Net Split Data:")
for key, value in net_split_data.items():
    doc.add_paragraph(f"{key}: {value}")



# Plotting
plt.bar(net_split_data.keys(), net_split_data.values())
plt.xlabel('Score')
plt.ylabel('Count')
plt.title('Net Split of All Reviews')

# Save the plot below the previously added data
plt.savefig(docx_file_path.replace('.docx', '_plot.png'))

# Add the plot to the Word document
doc.add_picture(docx_file_path.replace('.docx', '_plot.png'))

# Save the Word document
doc.save(docx_file_path)

















#######################################################################
# saving data in notepad


file_path = r"C:\Users\asus\Desktop\HDFC Work Related\text.txt"
data_to_write = f"Today's Date = {current_datetime.strftime('%Y-%m-%d')}\n"

with open(file_path, "a") as file:
    # Write data to the file
    file.write(data_to_write)
    
data_to_write=f"\n\n\nOver All Rating = {overall_rating}\nand Given interval Rating = {interval_rating} \n\n\n "

with open(file_path, "a") as file:
    # Write data to the file
    file.write(data_to_write)

# Write four lines of empty space
with open(file_path, "a") as file:
    file.write("\n" * 4)

# Write the DataFrame to the file in CSV format
df_xx.to_csv(file_path, mode='a', index=False, sep='\t')
 
#writing net split to note pad
# Write the dictionary data below the previously added data
with open(file_path, "a") as file:
    file.write("\writing net split Data:\n")
    for key, value in net_split_data.items():
        file.write(f"{key}: {value}\n")

plt.bar(net_split_data.keys(), net_split_data.values())
plt.xlabel('Score')
plt.ylabel('Count')
plt.title('Net Split of All Reviews')

# Save the plot below the previously added data
plt.savefig(file_path.replace('.txt', '_plot.png'))















en the text file using the default text editor
os.system(f'start {file_path}')














with open(file_path, "a") as file:
    # Write data to the file
    file.write(data_to_write)

# Write the DataFrame to a text file (CSV format)
df_xx.to_csv(file_path, index=False, sep='\t')  # Use tab as a separator for better readability in Notepad

#saving all the information on csv file




# DataFrame 1
data1 = {"Column1": [current_datetime]}
df1 = pd.DataFrame(data1)

# DataFrame 2
data2 = {"Column2": [overall_rating,interval_rating]}
df2 = pd.DataFrame(data2)

# DataFrame 2
data3 = {"Column3": [df_xx]}
df3 = pd.DataFrame(data3)
        

# Concatenate DataFrames vertically
result_df = pd.concat([df1, df2, df3], ignore_index=True)

result_df.to_csv('output_dataframe.csv', index=False)


    
    if T_minus_days=='all':
        rating_count_fun(full_filtered[0])
    if isinstance(time_duration, int):   
        rating_count_fun(full_filtered[1])
    return "wrong inputed data"
#########################################################################

    
    
    

for score in trend_line_data[0]:
        lis.append(score['score'])
        print(score['score'])
          
    
    while T_minus_days>=1:
        day_wise_reviews=[
            [trend_line_data[0][0]['score'],trend_line_data[0][0]['at']]
            for review in reviews_result
            if start_date <= review['at'] <= end_date
        ]
        
    
    
    return trend_line_data

    
trend_line_data=Trend_of_Tminus_days(T_minus_days)




