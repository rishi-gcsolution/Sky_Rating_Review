# -*- coding: utf-8 -*-
"""
Created on Fri Nov 17 10:35:18 2023

@author: asus
"""

# -*- coding: utf-8 -*-
"""
Created on Fri Nov 10 17:28:44 2023

@author: asus
"""

# -*- coding: utf-8 -*-
"""
Created on Fri Nov 10 11:35:03 2023

@author: asus
"""
import os
import pandas as pd
import requests
from datetime import datetime
from google_play_scraper import app, Sort, reviews
from datetime import datetime,timedelta, timezone
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Pt

# Replace 'com.cloudtradetech.sky' with the package name of the app you want to scrape.
app_id = 'com.cloudtradetech.sky'

# Specify language and country
language = 'en'
country = 'in'

 
# Specify start and end dates for filtering format date= yyyy-mm-dd
start_date_str = "2023-11-01"  # Replace with your desired start date
end_date_str = "2023-11-03"  # Replace with your desired end date
T_minus_days=5

##################################
# Current datetime module
##################################


def get_current_datetime():
    try:
        response = requests.get("http://worldtimeapi.org/api/ip")
        response.raise_for_status()  # Raise an exception for bad responses

        data = response.json()
        current_datetime_str = data['datetime']
        # Convert the datetime string to a datetime object
        current_datetime = datetime.strptime(current_datetime_str, "%Y-%m-%dT%H:%M:%S.%f%z")

        return current_datetime

    except requests.exceptions.RequestException as e:
        print(f"Error fetching current datetime: {e}")
        return None
   
def full_review(app_id=app_id, language=language, country=country):
    try:
        # Fetch the reviews sorted by date in descending order
        reviews_result, _ = reviews(
            app_id,
            lang=language,
            country=country,
            sort=Sort.NEWEST,
            count=10000,  # Number of reviews to retrieve (adjust as needed)
        )
        return reviews_result
    except Exception as e:
        print(f"Error fetching reviews: {e}")
        return None

def filtered_review(T_minus_days=T_minus_days, app_id=app_id, language=language, country=country):
    try:
        reviews_result = full_review(app_id, language, country=country)
        
        if reviews_result is None:
            # Handle the case where an error occurred during fetching reviews
            print("Unable to fetch reviews.")
            return None

        # Convert start and end dates to datetime objects
        current_datetime = get_current_datetime()
        start_date = current_datetime - timedelta(days=T_minus_days)
        end_date = current_datetime.replace(tzinfo=timezone.utc)

        # Make start_date offset-aware by adding timezone information
        start_date = start_date.replace(tzinfo=timezone.utc)

        # Filter reviews based on date range given as start date and end date
        filtered_reviews = [
            review
            for review in reviews_result
            if start_date <= review['at'].replace(tzinfo=timezone.utc) <= end_date
        ]
        return filtered_reviews
    except Exception as e:
        print(f"Error filtering reviews: {e}")
        return None

#qqqq=filtered_review("2023-11-11","2023-11-12")
#total_review=full_review()
 
################################
#calculation of overall rating and separate rating based on  datetime interval


overall_rating=0
interval_rating=0
for review in full_review():
    overall_rating+=review['score']
overall_rating=overall_rating/len(full_review())

interval_rating = sum([
    review['score']
    for review in filtered_review(T_minus_days)])/len(filtered_review(T_minus_days))



###################################
#Enter the no of T minus days
#T_minus_days=10

# Subtract days from the current datetime
#result_datetime = current_datetime - timedelta(days=T_minus_days)
#######################################################################################################

#######################################################################################################

def Trend_of_Tminus_days(T_minus_days, app_id=app_id, language=language, country=country):
    try:
        lis = []
        current_datetime = get_current_datetime()
        
        if current_datetime is None:
            # Handle the case where an error occurred during fetching current datetime
            print("Unable to fetch current datetime.")
            return None
        
        def scoreSum_and_date(required_date):
            day_reviews = full_review(app_id, language, country)
            day_sum = [score['score'] for score in day_reviews
                       if score['at'].strftime("%Y-%m-%d") == required_date
                       ]
            return [sum(day_sum), required_date, len(day_sum)]

        while T_minus_days > 0:
            required_date = (current_datetime - timedelta(days=T_minus_days)).strftime("%Y-%m-%d")
            lis.append(scoreSum_and_date(required_date))
            T_minus_days -= 1

        lis = [
            [round(sublist[0] / sublist[2], 2), sublist[1]] if sublist[2] != 0 else [0, sublist[1]]
            for sublist in lis
        ]

        return lis
    except Exception as e:
        print(f"Error in Trend_of_Tminus_days: {e}")
        return None

#######################################################
# Split based on total reviews given

#T_minus days is time duration
def net_split_of_all_reviews(duration_of_T_minusDays):
    net_split={'5':0,'4':0,'3':0,'2':0,'1':0}
    
    if duration_of_T_minusDays=="all":
        for score in full_review():
            net_split[str(score['score'])]+=1
            
    if isinstance(duration_of_T_minusDays, int):
        
        #end_date=get_current_datetime().strftime("%Y-%m-%d")
        #start_date=(get_current_datetime() - timedelta(days=duration_of_T_minusDays)).strftime("%Y-%m-%d")
                
        for score in filtered_review(T_minus_days):
            net_split[str(score['score'])]+=1
    
    return net_split
            
         
#########################################################################
#split trend wise with respect to previous days this function is giving no of 5star 4star 3star 2star 1star given perticular days
########################################################################
def split_trendwise_wrt_previous_days(duration_of_T_minusDays, app_id=app_id, language=language, country=country):
    try:
        rating_count = 0
        
        def rating_count_fun(data_input):
            # Create a DataFrame
            df = pd.DataFrame(data_input)
            # Converting datetime to string of format yyyy-mm-dd
            df['at'] = df['at'].dt.strftime("%Y-%m-%d")
            
            # Count the number of ratings for each score on each day
            rating_counts = df.groupby(['at', 'score']).size().unstack(fill_value=0)
            
            # Rename columns for clarity
            rating_counts.columns = [f'{col} Star Ratings' for col in rating_counts.columns]
            
            # Add a column for Total Ratings
            rating_counts['Total Ratings'] = rating_counts.sum(axis=1)
            
            # Reset the index for a clean DataFrame
            rating_counts.reset_index(inplace=True)
            
            return rating_counts

        if duration_of_T_minusDays == "all":
            rating_count = rating_count_fun(full_review(app_id, language, country))
        elif isinstance(duration_of_T_minusDays, int):
            #end_date = get_current_datetime().strftime("%Y-%m-%d")
            #start_date = (get_current_datetime() - timedelta(days=duration_of_T_minusDays)).strftime("%Y-%m-%d")
            rating_count = rating_count_fun(filtered_review(duration_of_T_minusDays))
        else:
            print("Invalid duration_of_T_minusDays. It should be 'all' or an integer.")
            return None

        return rating_count

    except Exception as e:
        print(f"Error in split_trendwise_wrt_previous_days: {e}")
        return None
    

    
def thumbs_count():
    pass
def get_no_Of_ratingsYesterday_avgRating(T_minus_days=T_minus_days, app_id=app_id, language=language, country=country):
    try:
        yesterdays = split_trendwise_wrt_previous_days(T_minus_days, app_id, language, country)
        # Print column names for debugging
        #print("Column Names in yesterdays DataFrame:", yesterdays.columns)

        if yesterdays is None:
            # Handle the case where an error occurred during trendwise rating count calculation
            print("Unable to fetch yesterday's ratings.")
            return None
        # Ensure all required columns are present
        required_columns = ['1 Star Ratings', '2 Star Ratings', '3 Star Ratings', '4 Star Ratings', '5 Star Ratings']
        for col in required_columns:
            if col not in yesterdays.columns:
                # Add the missing column with default value 0
                yesterdays[col] = 0
        # Calculate weighted average for each row across specified columns
        ratings_columns = required_columns
        weighted_average = (
            (yesterdays[ratings_columns] * [1, 2, 3, 4, 5]).sum(axis=1) / yesterdays['Total Ratings']
        )
        yesterdays['Weighted Average Rating'] = round(weighted_average, 2)

        return yesterdays

    except Exception as e:
        print(f"Error in get_no_Of_ratingsYesterday_avgRating: {e}")
        return None
def percent_split_change():
    
    
    pass
###########################################################################
# saving to ms Word

# Specify the file path for Word document
docx_file_path = r"C:\Users\asus\Desktop\HDFC Work Related\document_OnReviews.docx"


# Create a Word document
doc = Document()

# Write the date to the document
data_to_write = "                  Ratings with Reviews of HDFC SKY\n"
paragraph = doc.add_paragraph()

# Add runs for different parts of the string
paragraph.add_run(data_to_write[:data_to_write.find("Reviews")])  # Non-bold part
run_reviews = paragraph.add_run("Reviews")
run_reviews.bold = True  # Bold part
paragraph.add_run(data_to_write[data_to_write.find("Reviews") + len("Reviews"):])  # Non-bold part


'''
date_analysis = get_current_datetime().strftime('%Y-%m-%d')
#T_minus_days = 7  # Replace with your desired value
data_to_write = f" Date of Analysis = {date_analysis}\n also LAST {T_minus_days} days analysis \n"

# Add a paragraph to the document
doc.add_paragraph()
'''

yesterday=get_no_Of_ratingsYesterday_avgRating(T_minus_days)




# locating the latest date
yesterday['at'] = pd.to_datetime(yesterday['at'])
latest_date_row = yesterday.loc[yesterday['at'].idxmax()]


# Extract the Total Ratings from the latest date row
#total_ratings_latest_date = latest_date_row['Total Ratings']

data_to_write = f" Total no of reviews on {(get_current_datetime() - timedelta(days=1)).strftime('%dth %b')} =  {latest_date_row['Total Ratings']} \n Average ratings of all Reviews on {yesterday.loc[yesterday['at'].idxmax()][0].strftime('%dth %b')} = {round(overall_rating, 2)}\n"
doc.add_paragraph(data_to_write)


#data_to_write = "Rating Split Analysis"
#doc.add_paragraph(data_to_write)

# Write the DataFrame to the document in tabular form
doc.add_paragraph(f"Reviews analysis of Last {T_minus_days} days:")

#changing the order of yesterday dataframe

#changing the column order
# Specify the desired column order
desired_columns = ['at','Total Ratings', '5 Star Ratings', '4 Star Ratings', '3 Star Ratings', '2 Star Ratings', '1 Star Ratings', 'Weighted Average Rating']

# Create a new DataFrame with the reordered columns
yesterday = yesterday[desired_columns]

yesterday['at']=yesterday['at'].dt.strftime('%b %d')
# Rename the 'A' column to 'X'
yesterday.rename(columns={'Weighted Average Rating': 'Average Rating'}, inplace=True)


table = doc.add_table(yesterday.shape[0]+1, yesterday.shape[1])

for col_num, col_name in enumerate(yesterday.columns):
    table.cell(0, col_num).text = col_name
    for row_num in range(yesterday.shape[0]):
        table.cell(row_num+1, col_num).text = str(yesterday.iloc[row_num, col_num])

# Write two lines of empty space
doc.add_paragraph("\n" * 1)


'''
# Write the DataFrame to the document in tabular form
doc.add_paragraph(f"Trend of Last {T_minus_days} selected days ")
xx=Trend_of_Tminus_days(T_minus_days)
# Convert list of lists to DataFrame
df_xx = pd.DataFrame(xx, columns=['avg daywise rating', 'Date'])
df_xx = df_xx[['Date', 'avg daywise rating']]
table = doc.add_table(df_xx.shape[0]+1, df_xx.shape[1])
for col_num, col_name in enumerate(df_xx.columns):
    table.cell(0, col_num).text = col_name
    for row_num in range(df_xx.shape[0]):
        table.cell(row_num+1, col_num).text = str(df_xx.iloc[row_num, col_num])

'''


# Extract x and y values from the list
#xx=xx[::-1]
xx=Trend_of_Tminus_days(T_minus_days)
x_values = [pair[1] for pair in xx]
y_values = [pair[0] for pair in xx]

# Convert date strings to datetime objects
date_objects = [datetime.strptime(date_str, '%Y-%m-%d') for date_str in x_values]

# Write two lines of empty space
#doc.add_paragraph("\n" * 2)

# Write net split data of overall and selected LAST no of days to the document
doc.add_paragraph("Net Rating Distribution based on All reviews given")
#doc.add_paragraph("Net Split Data of overall rating \n\n")
# Write net split data to the document
net_split_data = net_split_of_all_reviews('all')
doc.add_paragraph("Net Split Data:")
for key, value in net_split_data.items():
    doc.add_paragraph(f"{key}: {value}")

'''
#trend of last selected days DATA and LINE GRAPH
# Plotting
plt.bar(net_split_data.keys(), net_split_data.values())
plt.xlabel('Rating Score')
plt.ylabel('Count')
plt.title('Net Rating Distribution based on All reviews given')

# Save the plot below the previously added data
plt.savefig(docx_file_path.replace('.docx', '_plot_netSplit.png'), dpi=50)

# Close the plot to free up resources
plt.close()

# Add the plot to the Word document
doc.add_picture(docx_file_path.replace('.docx', '_plot_netSplit.png'))
'''




# Plotting both line graph and bar graph on the same plot# Plotting both line graph and bar graph on the same plot
fig, ax1 = plt.subplots()

# Bar graph (Total Rating)
bar_plot = ax1.bar(yesterday['at'], yesterday['Total Ratings'], color='b', alpha=0.7, label='Total Rating')

# Display values at points on the bar graph
for bar, value in zip(bar_plot, yesterday['Total Ratings']):
    height = bar.get_height()
    ax1.text(bar.get_x() + bar.get_width() / 2, height, value, ha='center', va='bottom', color='black')

#ax1.set_xlabel('Date')
ax1.set_ylabel('Total Rating', color='b')
ax1.tick_params('y', colors='b')

# Create a second y-axis for the line graph
ax2 = ax1.twinx()

# Line graph (Weighted Average Rating)
#either use x_values or same date value can also be extracted via yesterday variable 
line_plot = ax2.plot(yesterday['at'], yesterday['Average Rating'], linestyle='-', color='r', label='Average Rating')


# Display values at points on the line graph
for point, value in zip(line_plot[0].get_data()[0], yesterday['Average Rating']):
    ax2.text(point, value, value, ha='right', va='bottom', color='black')

ax2.set_ylabel('Avg Rating', color='r')
ax2.tick_params('y', colors='r')

# Set title and legend
#plt.title('Line and Bar Graphs')
fig.tight_layout()
plt.legend(loc='upper left')

# Save the plot below the previously added data
plt.savefig(docx_file_path.replace('.docx', '_plot_netSplit.png'))

# Close the plot to free up resources
plt.close()

# Add the plot to the Word document
doc.add_picture(docx_file_path.replace('.docx', '_plot_netSplit.png'))



'''

# Write two lines of empty space
doc.add_paragraph("\n" * 1)

xx=Trend_of_Tminus_days(T_minus_days)
# Assuming xx is a list of pairs [(rating, date_str), ...]
x_values = [pair[1] for pair in xx]
y_values = [pair[0] for pair in xx]

# Convert date strings to datetime objects
date_objects = [datetime.strptime(date_str, '%Y-%m-%d') for date_str in x_values]

# Format datetime objects to 'Mon Day' format
x_values = [date_obj.strftime('%b %d') for date_obj in date_objects]

# Plot the line plot
plt.figure(figsize=(7,4))

# Subplot 1: Rating Trend line Analysis
plt.subplot(1, 2, 1)
plt.plot(x_values, y_values, marker='o', linestyle='-', color='b')
plt.xlabel('Date')
plt.ylabel('Rating')
plt.title(f'Trend line of LAST {T_minus_days} days')
plt.grid(True)


# Display values on the line plot
for i, txt in enumerate(y_values):
    plt.annotate(txt, (x_values[i], y_values[i]))

# Subplot 2: Rating Distribution based on selected Last {T_minus_days} days
plt.subplot(1, 2, 2)

# Assuming net_split_data is a dictionary with rating scores as keys and counts as values
plt.bar(net_split_data.keys(), net_split_data.values())
plt.xlabel('Rating Score')
plt.ylabel('Count')
plt.title(f'Rating Distribution of {T_minus_days} days')

# Display values on the bar plot
for key, value in net_split_data.items():
    plt.text(key, value, str(value), ha='center', va='bottom')
    


# Save the combined plot
plt.savefig(docx_file_path.replace('.docx', '_combined_plots.png'))

# Close the plot to free up resources
plt.close()

# Add the combined plot to the Word document
doc.add_picture(docx_file_path.replace('.docx', '_combined_plots.png'))

'''
# Write two lines of empty space
doc.add_paragraph("\n" * 1)

#displaying all review details of last selected days

# Write the DataFrame to the document in tabular form
doc.add_paragraph(f"Displaying all the reviews of last {T_minus_days-3} Days along with all related details\n")
end_date=get_current_datetime().strftime("%Y-%m-%d")
start_date=(get_current_datetime() - timedelta(days=T_minus_days)).strftime("%Y-%m-%d")
yesterday=filtered_review(T_minus_days-2)
yesterday=pd.DataFrame(yesterday)
# Extract specific columns
selected_columns = ['userName', 'content', 'at', 'score', 'thumbsUpCount', 'appVersion', 'replyContent']
yesterday = yesterday[selected_columns]
# Rename the 'A' column to 'X'
yesterday.rename(columns={'userName': 'Name'}, inplace=True)
yesterday.rename(columns={'content': 'Review'}, inplace=True)
yesterday.rename(columns={'score': 'Rating'}, inplace=True)

# Convert 'at' column to datetime object
yesterday['at'] = pd.to_datetime(yesterday['at'])

# Format 'at' column without seconds
yesterday['at'] = yesterday['at'].dt.strftime('%H:%M')
#changing the name of the column at to date
yesterday.rename(columns={'at': 'Time'}, inplace=True)


table = doc.add_table(yesterday.shape[0]+1, yesterday.shape[1])
for col_num, col_name in enumerate(yesterday.columns):
    table.cell(0, col_num).text = col_name
    for row_num in range(yesterday.shape[0]):
        table.cell(row_num+1, col_num).text = str(yesterday.iloc[row_num, col_num])

# Write two lines of empty space
doc.add_paragraph("\n" * 1)



# Save the Word document
doc.save(docx_file_path)
