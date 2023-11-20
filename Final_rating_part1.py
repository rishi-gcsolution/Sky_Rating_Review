# -*- coding: utf-8 -*-
"""
Created on Mon Nov 13 20:14:26 2023

@author: asus
"""


import google_play_scraper
import urllib.request
from bs4 import BeautifulSoup
import os
import pandas as pd
import requests
from datetime import datetime
from google_play_scraper import app, Sort, reviews
from datetime import datetime,timedelta
import matplotlib.pyplot as plt
from docx import Document


# Extract the latest x_days of data for each column
x_days=5

def get_rating_via_webGoogleapp(url):
#Send a GET request to the URL
    try:
        response = urllib.request.urlopen(url)
    
       # Check if the request was successful (status code 200)
        if response.getcode() == 200:
            # Parse the HTML content of the page
            soup = BeautifulSoup(response.read(), 'html.parser')
            
            # Extract and print text from all div elements with class "TT9eCd"
            for row in soup.find_all('div', attrs={"class": "TT9eCd"}):
                print(row.text)
        
            # Find the div element with itemprop="starRating" and class "TT9eCd"
            rating_element = soup.find('div', class_='TT9eCd')
            
            # Check if the element is found
            if rating_element:
                # Extract the text content within the element
                rating = rating_element.text.strip()
                print("Rating:", rating)
        return rating
    except Exception as e:
        print(f"Error during web scraping: {e}")
        return None


def get_current_datetime():
    try:
        response = requests.get("http://worldtimeapi.org/api/ip")
        response.raise_for_status()  # Raise an exception for bad responses
        data = response.json()
        current_datetime_str = data['datetime']
        # Convert the datetime string to a datetime object
        current_datetime = datetime.strptime(current_datetime_str, "%Y-%m-%dT%H:%M:%S.%f%z")

        return current_datetime

    except requests.exceptions.RequestException as e:
        print(f"Error fetching current datetime: {e}")
        return None

def get_app_rating(app_id):
    try:
        app_info = google_play_scraper.app(app_id)
        rating = app_info['score']
        return round(rating,2)
    except Exception as e:
        print(f"Error: {e}")

# Option 2: Create a DataFrame from a dictionary
data_dict = { 'HDFC SKY': [get_rating_via_webGoogleapp("https://play.google.com/store/apps/details?id=com.cloudtradetech.sky&hl=en_US")[:-4]],
              'HDFC INVESTRIGHT': [get_app_rating('com.hsl.investright')],
              'GROWW': [get_app_rating('com.nextbillion.groww')],
              'ZERODHA': [get_app_rating('com.zerodha.kite3')],
              'PAYTMMONEY': [get_app_rating('com.paytmmoney')],
              'ANGLE ONE': [get_app_rating('com.msf.angelmobile')],
              'DHAN': [get_app_rating('com.dhan.live')],
              'UPSTOX': [get_app_rating('in.upstox.app')]
             }
df_Google_Ratings = pd.DataFrame(data_dict)

#Replace 'your_file.csv' with the actual file path
file_path_csv = r'C:\Users\asus\Desktop\HDFC Work Related\excel_files\Daily rating distribution _ 2023-10-15 - 2023-11-11.csv'

# Read the CSV file into a DataFrame
df_csv = pd.read_csv(file_path_csv)


# Calculate the sum of star ratings column-wise(total Rating)
df_csv['Total Rating'] = df_csv[['5 stars', '4 stars', '3 stars', '2 stars', '1 star']].sum(axis=1)
total_rating_sum = df_csv['Total Rating'].sum()

# Breakup of Ratings of LAST 5 DAYS 


# Convert the 'Date' column to datetime
df_csv['Date'] = pd.to_datetime(df_csv['Date'])

# Sort the DataFrame based on the 'Date' column
df_csv_sorted = df_csv.sort_values(by='Date', ascending=False)

# Filter the latest x rows
latest_rows = df_csv_sorted.head(x_days)




def get_weighted_avg():
               
    # Calculate weighted average for each row across specified columns
    ratings_columns = ['5 stars', '4 stars', '3 stars', '2 stars', '1 star']
    weighted_average = (
        (latest_rows[ratings_columns] * [5,4,3,2,1]).sum(axis=1) / latest_rows['Total Rating']
    )
    latest_rows['Weighted Average Rating'] = round(weighted_average,2)
    return latest_rows

latest_rows=get_weighted_avg()
latest_rows['Date'] = latest_rows['Date'].dt.strftime('%b %d')
#resetting the index
latest_rows=latest_rows.reset_index(drop=True)

# Drop a certain column by name
column_to_drop = 'Notes'
latest_rows = latest_rows.drop(column_to_drop, axis=1)
#changing the column order
# Specify the desired column order
desired_columns = ['Date', 'Total Rating', '5 stars', '4 stars', '3 stars', '2 stars', '1 star', 'Weighted Average Rating']

# Create a new DataFrame with the reordered columns
latest_rows = latest_rows[desired_columns]

##################################################################################
# Create a Word document

# saving to ms Word

# Specify the file path for Word document
docx_file_path = r"C:\Users\asus\Desktop\HDFC Work Related\document.docx"
doc = Document()

# Write the date to the document


data_to_write = f" {get_current_datetime().strftime('%dth %b')}\n"
doc.add_paragraph(data_to_write)

data_to_write = "                       HDFC SKY ANDROID APP\n"
doc.add_paragraph(data_to_write)

data_to_write = f" {(get_current_datetime() - timedelta(days=1)).strftime('%dth %b')} -- Rating and Reviews\n"
doc.add_paragraph(data_to_write)

data_to_write = f" Google Rating as of {(get_current_datetime() - timedelta(days=1)).strftime('%dth %b')}\n"
doc.add_paragraph(data_to_write)

# printing Google Rating Sections
# Write the DataFrame to the document in tabular form

table = doc.add_table(df_Google_Ratings.shape[0]+1, df_Google_Ratings.shape[1])

for col_num, col_name in enumerate(df_Google_Ratings.columns):
    table.cell(0, col_num).text = col_name
    for row_num in range(df_Google_Ratings.shape[0]):
        table.cell(row_num+1, col_num).text = str(df_Google_Ratings.iloc[row_num, col_num])

# Write two lines of empty space
doc.add_paragraph("\n" * 2)


data_to_write = f" Total no of rating on {latest_rows.loc[0, 'Date']} = {latest_rows.loc[0, 'Total Rating']} "
doc.add_paragraph(data_to_write)


data_to_write = f"Average of Ratings on {latest_rows.loc[0, 'Date']} = {latest_rows.loc[0, 'Weighted Average Rating']} "
doc.add_paragraph(data_to_write)


# Write the DataFrame to the document in tabular form
doc.add_paragraph(f"Break up of ratings of {x_days} days")

latest_rows.rename(columns={'Weighted Average Rating': 'Average Rating'}, inplace=True)

table = doc.add_table(latest_rows.shape[0]+1, latest_rows.shape[1])

for col_num, col_name in enumerate(latest_rows.columns):
    table.cell(0, col_num).text = col_name
    for row_num in range(latest_rows.shape[0]):
        table.cell(row_num+1, col_num).text = str(latest_rows.iloc[row_num, col_num])

# Write two lines of empty space
doc.add_paragraph("\n" * 11)

doc.add_paragraph(f'Trend of last {x_days} (Average Rating V/S Total Rating)')


# Plotting both line graph and bar graph on the same plot# Plotting both line graph and bar graph on the same plot
fig, ax1 = plt.subplots()

# Bar graph (Total Rating)
bar_plot = ax1.bar(latest_rows['Date'][::-1], latest_rows['Total Rating'][::-1], color='b', alpha=0.7, label='Total Rating')

# Display values at points on the bar graph
for bar, value in zip(bar_plot, latest_rows['Total Rating'][::-1]):
    height = bar.get_height()
    ax1.text(bar.get_x() + bar.get_width() / 2, height, value, ha='center', va='bottom', color='black')

#ax1.set_xlabel('Date')
ax1.set_ylabel('Total Rating', color='b')
ax1.tick_params('y', colors='b')

# Create a second y-axis for the line graph
ax2 = ax1.twinx()

# Line graph (Weighted Average Rating)
line_plot = ax2.plot(latest_rows['Date'][::-1], latest_rows['Average Rating'][::-1], marker='o', linestyle='-', color='r', label='Average Rating',alpha=0.7)

# Display values at points on the line graph
for point, value in zip(line_plot[0].get_data()[0], latest_rows['Average Rating'][::-1]):
    ax2.text(point, value, value, ha='right', va='bottom', color='black')

ax2.set_ylabel('Avg Rating', color='r')
ax2.tick_params('y', colors='r')

# Set title and legend
#plt.title('Line and Bar Graphs')
fig.tight_layout()
plt.legend(loc='upper left')

# Save the plot below the previously added data
plt.savefig(docx_file_path.replace('.docx', '_plot_netSplit.png'))

# Close the plot to free up resources
plt.close()

# Add the plot to the Word document
doc.add_picture(docx_file_path.replace('.docx', '_plot_netSplit.png'))
# Save the Word document
doc.save(docx_file_path)




